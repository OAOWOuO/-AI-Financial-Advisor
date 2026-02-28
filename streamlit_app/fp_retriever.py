"""
fp_retriever.py — Document ingestion and RAG retrieval for the Financial Planning Assistant.

Architecture: in-memory numpy cosine similarity (no external vector DB required).
Documents are stored in st.session_state across a browser session.
Each chunk carries rich metadata: source_type, topic, reliability_level, date.

Supports: PDF, Markdown, TXT, HTML (text portion), Word (.docx).
"""
from __future__ import annotations
import io, os, hashlib
from typing import List, Dict, Any, Tuple, Optional

import streamlit as st
import numpy as np


# ── Constants ─────────────────────────────────────────────────────────────────

EMBED_MODEL  = "text-embedding-3-small"
CHUNK_SIZE   = 400    # words per chunk
CHUNK_OVERLAP = 50

TOPIC_TAGS  = ["retirement", "tax", "debt", "insurance", "cash_flow",
               "estate", "investing", "goals", "general"]
SOURCE_TYPES = ["official_guidance", "case_study", "worksheet", "article",
                "class_material", "reference"]
RELIABILITY  = ["official", "educational", "secondary"]


# ── Session-level in-memory store ────────────────────────────────────────────

def _get_store() -> Dict[str, list]:
    if "fp_doc_store" not in st.session_state:
        st.session_state.fp_doc_store = {
            "ids": [], "documents": [], "embeddings": [], "metadatas": []
        }
    return st.session_state.fp_doc_store


def store_count() -> int:
    return len(_get_store()["ids"])


def list_sources() -> List[Dict[str, Any]]:
    """Return one row per unique source document (filename)."""
    seen: Dict[str, Dict] = {}
    for m in _get_store()["metadatas"]:
        src = m.get("source", "?")
        if src not in seen:
            seen[src] = {
                "source":       src,
                "topic":        m.get("topic", "general"),
                "source_type":  m.get("source_type", "reference"),
                "reliability":  m.get("reliability_level", "educational"),
                "date":         m.get("date", ""),
                "chunks":       0,
            }
        seen[src]["chunks"] += 1
    return list(seen.values())


def clear_store() -> None:
    if "fp_doc_store" in st.session_state:
        del st.session_state.fp_doc_store


# ── Chunking ──────────────────────────────────────────────────────────────────

def _chunk_text(text: str, source: str, page: int,
                metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
    words = text.split()
    chunks: List[Dict[str, Any]] = []
    i = 0
    while i < len(words):
        window = words[i: i + CHUNK_SIZE]
        chunk_id = hashlib.md5(f"{source}|{page}|{i}".encode()).hexdigest()[:16]
        chunks.append({
            "text":     " ".join(window),
            "chunk_id": chunk_id,
            "metadata": {**metadata, "source": source, "page": page, "chunk_id": chunk_id},
        })
        if len(window) < CHUNK_SIZE:
            break
        i += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


# ── Ingestion ──────────────────────────────────────────────────────────────────

def ingest_bytes(
    file_bytes: bytes,
    filename: str,
    openai_client,
    topic: str        = "general",
    source_type: str  = "reference",
    reliability: str  = "educational",
    date: str         = "",
) -> Tuple[int, str]:
    """
    Ingest a document from raw bytes.
    Returns (chunks_added, error_message).  error_message is "" on success.
    """
    base_meta = {
        "topic":            topic,
        "source_type":      source_type,
        "reliability_level": reliability,
        "date":             date,
    }

    all_chunks: List[Dict[str, Any]] = []
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    all_chunks.extend(_chunk_text(text, filename, page_num, base_meta))
        except Exception as e:
            return 0, f"Failed to read PDF: {e}"

    elif ext in ("md", "txt"):
        text = file_bytes.decode("utf-8", errors="replace")
        all_chunks.extend(_chunk_text(text, filename, 0, base_meta))

    elif ext == "html":
        try:
            import re
            text = file_bytes.decode("utf-8", errors="replace")
            text = re.sub(r"<[^>]+>", " ", text)   # strip HTML tags
            text = re.sub(r"\s+", " ", text).strip()
            all_chunks.extend(_chunk_text(text, filename, 0, base_meta))
        except Exception as e:
            return 0, f"Failed to parse HTML: {e}"

    elif ext == "doc":
        return 0, (
            "Old .doc format (Word 97-2003) is not supported. "
            "Please re-save the file as .docx or export it as PDF, then upload again."
        )

    elif ext == "docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_bytes))
            # Group paragraphs into logical "pages" (~50 paragraphs each for chunking)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # Also pull text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            if not paragraphs:
                return 0, "No text found in this Word document."
            page_size = 50   # paragraphs per virtual page
            for page_num, start in enumerate(range(0, len(paragraphs), page_size), start=1):
                page_text = "\n".join(paragraphs[start: start + page_size])
                all_chunks.extend(_chunk_text(page_text, filename, page_num, base_meta))
        except ImportError:
            return 0, "python-docx is not installed. Run: pip install python-docx"
        except Exception as e:
            return 0, f"Failed to read Word document: {e}"

    else:
        return 0, f"Unsupported file type '.{ext}'. Supported: pdf, docx, md, txt, html."

    if not all_chunks:
        return 0, "No text could be extracted from this file."

    # Filter already-indexed chunks
    store = _get_store()
    existing_ids = set(store["ids"])
    new_chunks = [c for c in all_chunks if c["chunk_id"] not in existing_ids]
    if not new_chunks:
        return 0, ""

    # Embed in batches of 100
    total_added = 0
    batch_size  = 100
    for i in range(0, len(new_chunks), batch_size):
        batch  = new_chunks[i: i + batch_size]
        texts  = [c["text"] for c in batch]
        try:
            resp       = openai_client.embeddings.create(model=EMBED_MODEL, input=texts)
            embeddings = [r.embedding for r in resp.data]
        except Exception as e:
            return total_added, f"Embedding API error: {e}"

        for j, chunk in enumerate(batch):
            if chunk["chunk_id"] not in existing_ids:
                store["ids"].append(chunk["chunk_id"])
                store["documents"].append(chunk["text"])
                store["embeddings"].append(embeddings[j])
                store["metadatas"].append(chunk["metadata"])
                existing_ids.add(chunk["chunk_id"])
                total_added += 1

    return total_added, ""


# ── Retrieval ─────────────────────────────────────────────────────────────────

def retrieve(
    query: str,
    openai_client,
    top_k: int     = 8,
    topic_filter: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """
    Retrieve top-k relevant chunks for a query.
    Optionally filter by topic tag.
    Returns list of {text, metadata, score}.
    """
    store = _get_store()
    if not store["embeddings"]:
        return []

    # Embed query
    resp = openai_client.embeddings.create(model=EMBED_MODEL, input=query)
    q    = np.array(resp.data[0].embedding, dtype="float32")
    q   /= np.linalg.norm(q) + 1e-10

    # Filter by topic if requested
    indices = range(len(store["embeddings"]))
    if topic_filter and topic_filter != "all":
        indices = [i for i in indices
                   if store["metadatas"][i].get("topic", "") == topic_filter]

    if not indices:
        return []

    scores = []
    for i in indices:
        e = np.array(store["embeddings"][i], dtype="float32")
        e /= np.linalg.norm(e) + 1e-10
        scores.append((float(q @ e), i))

    scores.sort(reverse=True)
    top = scores[:top_k]

    return [
        {
            "text":     store["documents"][i],
            "metadata": store["metadatas"][i],
            "score":    round(s, 4),
        }
        for s, i in top
    ]
